import os
import fitz  #used to open and work with PDF files
import docx  #python-docx – used for handling Word docs
import pytesseract  #OCR to extract text from images
import shutil
import json
import spacy  #NLP for parsing text
import requests  #for downloading files from the web
import re
import pandas as pd
import glob
from PIL import Image  #image handling
from docx.api import Document
import openai
from pdf2image import convert_from_path  #fallback to screenshot whole PDF pages

#load small english NLP model for text processing
nlp = spacy.load("en_core_web_sm")

#API setup
OPENAI_API_KEY = ""
if not OPENAI_API_KEY:
    raise ValueError("OpenAI API key not found. Set it as an environment variable.")

client = openai.OpenAI(api_key=OPENAI_API_KEY)

#defining folder paths
INPUT_FOLDER = "./input_reports"
OUTPUT_FOLDER = "./output_visuals"
DOWNLOAD_FOLDER = "./downloaded_reports"

#make sure the folders exist
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

#keywords used to extract datapoints for text. Bascially anytime one of these shows up I extract the entire sentence. 
#I'm don't think this is a great system, also not sure what key words to include. BRAINSTORM
KEY_TERMS = [
    "production", "capacity", "sales value", "exports", "imports", "shipments", "consumption",
    "employment", "PPI", "scrap", "BOF", "EAF", "casting", "mill", "decarbonization", "DOE",
    "price", "value", "tons", "million tons", "net import", "market share", "funding", "steel",
    "slag", "throughput", "billion", "subsidy", "grant", "CO2", "energy", "labor"
]

#use GPT to generate descriptions + tags from image OCR and surrounding context
def generate_metadata(caption, ocr_text, context, source_url):
    prompt = f"""
The following is an excerpt from a report and OCR text from an associated figure. Based on this information, generate:
1. A 1-2 sentence description of the figure
2. A list of 5-10 tags for later search and classification

OCR Text:
{ocr_text}

Caption (if available):
{caption}

Surrounding Text:
{context}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    content = response.choices[0].message.content
    lines = content.strip().split("\n")
    description = ""
    tags = []
    for line in lines:
        if line.strip().startswith("1."):
            description = line.split("1.")[-1].strip()
        elif line.strip().startswith("2."):
            tag_line = line.split("2.")[-1].strip()
            tags = [tag.strip() for tag in re.split(r",|\n|\||-", tag_line) if tag.strip()]
        elif not line.startswith("1.") and not line.startswith("2.") and description:
            description += " " + line.strip()
    return description, tags

#break full text into sentences and pull out anything that matches key terms
def extract_datapoints_from_text(text, source_url):
    sentences = re.split(r'(?<=[.!?]) +', text)
    datapoints = []
    for sentence in sentences:
        if any(term.lower() in sentence.lower() for term in KEY_TERMS):
            if ";" in sentence or " and " in sentence:
                sub_sentences = re.split(r'[;.]', sentence)
                for sub in sub_sentences:
                    if any(term.lower() in sub.lower() for term in KEY_TERMS):
                        datapoints.append({"datapoint_text": sub.strip(), "source_url": source_url})
            else:
                datapoints.append({"datapoint_text": sentence.strip(), "source_url": source_url})
    return datapoints

#download PDF or DOCX and save locally -- I have a folder with downloaded reports. 
def download_file(url, save_dir):
    filename = url.split("/")[-1]
    filepath = os.path.join(save_dir, filename)
    r = requests.get(url)
    with open(filepath, 'wb') as f:
        f.write(r.content)
    return filepath, url

#this function processes PDF files (main logic)
def process_pdf(file_path, source_url):
    doc = fitz.open(file_path) #open the PDF
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_dir = os.path.join(OUTPUT_FOLDER, base_name)
    os.makedirs(output_dir, exist_ok=True)

    for page_number in range(len(doc)):
        page = doc[page_number]
        images = page.get_images(full=True) #this part filters through the page to check for images

        extracted_any = False  #track whether any figure was processed

        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_filename = f"{base_name}_page{page_number+1}_fig{img_index+1}.{image_ext}"
            image_path = os.path.join(output_dir, image_filename)

            with open(image_path, "wb") as f:
                f.write(image_bytes)

            img_pil = Image.open(image_path)
            ocr_text = pytesseract.image_to_string(img_pil) #try to pull text from the image

            if len(ocr_text.split()) < 5: #filter out images with less than 5 words, bc they're probably logos or unhelpful images
                os.remove(image_path)  #discard logo or decorative images
                continue

            text_blocks = page.get_text("blocks")
            text_blocks.sort(key=lambda b: b[1])
            all_text = "\n".join([b[4] for b in text_blocks]) #grab all text from the page

            #look for a "Figure" label as a rough caption
            caption = ""
            for b in text_blocks:
                if "Figure" in b[4]:
                    caption = b[4]
                    break

            #use GPT to get description/tags
            description, tags = generate_metadata(caption, ocr_text, all_text, source_url)

            metadata = {
                "filename": image_filename,
                "caption": caption,
                "ocr_text": ocr_text,
                "page_number": page_number + 1,
                "source_file": file_path,
                "source_link": source_url,
                "description": description,
                "tags": tags
            }

            #write out JSON metadata for this figure
            with open(os.path.join(output_dir, image_filename + ".json"), "w") as f:
                json.dump(metadata, f, indent=2)

            extracted_any = True  # flag we processed at least one figure

        #Fallback-- Screenshot entire page if no valid images were extracted
        if not extracted_any:
            pil_images = convert_from_path(file_path, first_page=page_number+1, last_page=page_number+1)
            for i, pil_img in enumerate(pil_images):
                image_filename = f"{base_name}_page{page_number+1}_screenshot{i+1}.png"
                image_path = os.path.join(output_dir, image_filename)
                pil_img.save(image_path)

                ocr_text = pytesseract.image_to_string(pil_img)
                if len(ocr_text.split()) < 5:
                    os.remove(image_path)  #skip blank pages
                    continue

                all_text = page.get_text()
                description, tags = generate_metadata("", ocr_text, all_text, source_url)
                metadata = {
                    "filename": image_filename,
                    "caption": "",
                    "ocr_text": ocr_text,
                    "page_number": page_number + 1,
                    "source_file": file_path,
                    "source_link": source_url,
                    "description": description,
                    "tags": tags
                }

                with open(os.path.join(output_dir, image_filename + ".json"), "w") as f:
                    json.dump(metadata, f, indent=2)

        #always extract datapoints from text even if no images
        all_text = page.get_text()
        datapoints = extract_datapoints_from_text(all_text, source_url)
        for dp in datapoints:
            dp["page_number"] = page_number + 1
        if datapoints:
            datapoints_output_path = os.path.join(output_dir, f"{base_name}_page{page_number+1}_datapoints.json")
            with open(datapoints_output_path, "w") as f:
                json.dump(datapoints, f, indent=2)

#same thing but for Word docs
def process_docx(file_path, source_url):
    doc = Document(file_path)
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_dir = os.path.join(OUTPUT_FOLDER, base_name)
    os.makedirs(output_dir, exist_ok=True)

    rels = doc.part._rels
    img_index = 0

    for rel in rels:
        rel = rels[rel]
        if "image" in rel.target_ref:
            img_index += 1
            img_data = rel.target_part.blob
            img_ext = rel.target_ref.split(".")[-1]
            img_filename = f"{base_name}_img{img_index}.{img_ext}"
            img_path = os.path.join(output_dir, img_filename)

            with open(img_path, "wb") as f:
                f.write(img_data)

            img = Image.open(img_path)
            ocr_text = pytesseract.image_to_string(img)
            word_count = len(ocr_text.strip().split())
            if word_count < 3:
                os.remove(img_path)
                continue

            context = "\n".join([para.text for para in doc.paragraphs])
            description, tags = generate_metadata("", ocr_text, context, source_url)

            metadata = {
                "filename": img_filename,
                "caption": "",
                "ocr_text": ocr_text,
                "page_number": None,
                "source_file": file_path,
                "source_link": source_url,
                "description": description,
                "tags": tags
            }

            with open(os.path.join(output_dir, img_filename + ".json"), "w") as f:
                json.dump(metadata, f, indent=2)

    context = "\n".join([para.text for para in doc.paragraphs])
    datapoints = extract_datapoints_from_text(context, source_url)
    if datapoints:
        datapoints_output_path = os.path.join(output_dir, f"{base_name}_datapoints.json")
        with open(datapoints_output_path, "w") as f:
            json.dump(datapoints, f, indent=2)


#ADD SOURCES HERE
URLS = [
    "https://www.energy-transitions.org/wp-content/uploads/2023/03/MPP-Breakthrough-Steel-US-v673.pdf"
    # , "Source 2", "Source 3", ...
]

#loop through each URL and process it based on file type
for url in URLS:
    local_path, source_url = download_file(url, DOWNLOAD_FOLDER)
    if local_path.endswith(".pdf"):
        process_pdf(local_path, source_url)
    elif local_path.endswith(".docx"):
        process_docx(local_path, source_url)
